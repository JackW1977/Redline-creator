"""Create test documents for validating the revision comparison system.

Generates two .docx files:
- test_early_rev.docx: The "early revision" with comments
- test_latest_rev.docx: The "latest revision" with edits but no comments

This uses python-docx for test document creation only.
Install: pip install python-docx
"""

from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime


def add_comment(doc, paragraph, text, author="Test Reviewer", initials="TR",
                anchor_start=0, anchor_end=None):
    """Add a comment to a paragraph in a python-docx document.

    This manipulates the XML directly to add proper Word comments.
    """
    # Get or create comments part
    if not hasattr(doc, '_comment_id'):
        doc._comment_id = 0

    comment_id = doc._comment_id
    doc._comment_id += 1

    # Ensure the document has a comments part
    from docx.opc.constants import RELATIONSHIP_TYPE as RT
    try:
        comments_part = doc.part.package.part_related_by(
            "http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments"
        )
    except KeyError:
        comments_part = None

    if comments_part is None:
        # We need to add comments via XML manipulation
        _add_comment_xml(doc, paragraph, comment_id, text, author, initials)
    else:
        _add_comment_xml(doc, paragraph, comment_id, text, author, initials)

    return comment_id


def _add_comment_xml(doc, paragraph, comment_id, text, author, initials):
    """Add a comment using direct XML manipulation."""
    W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    W14 = "http://schemas.microsoft.com/office/word/2010/wordml"

    # Add commentRangeStart at beginning of paragraph
    range_start = OxmlElement("w:commentRangeStart")
    range_start.set(qn("w:id"), str(comment_id))

    # Add commentRangeEnd at end of paragraph
    range_end = OxmlElement("w:commentRangeEnd")
    range_end.set(qn("w:id"), str(comment_id))

    # Add comment reference run
    ref_run = OxmlElement("w:r")
    ref_rpr = OxmlElement("w:rPr")
    ref_style = OxmlElement("w:rStyle")
    ref_style.set(qn("w:val"), "CommentReference")
    ref_rpr.append(ref_style)
    ref_run.append(ref_rpr)
    ref_elem = OxmlElement("w:commentReference")
    ref_elem.set(qn("w:id"), str(comment_id))
    ref_run.append(ref_elem)

    # Insert into paragraph XML
    p_elem = paragraph._p
    runs = list(p_elem.iter(qn("w:r")))

    if runs:
        runs[0].addprevious(range_start)
        runs[-1].addnext(range_end)
        range_end.addnext(ref_run)
    else:
        p_elem.append(range_start)
        p_elem.append(range_end)
        p_elem.append(ref_run)

    # Now add the comment definition to comments.xml
    # We need to create/update the comments part
    _ensure_comments_part(doc, comment_id, text, author, initials)


def _ensure_comments_part(doc, comment_id, text, author, initials):
    """Ensure the comments.xml part exists and add the comment to it."""
    from docx.opc.part import Part
    from docx.opc.constants import RELATIONSHIP_TYPE as RT

    W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    W14 = "http://schemas.microsoft.com/office/word/2010/wordml"

    comments_rel_type = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments"

    # Check if comments part already exists
    comments_xml = None
    for rel in doc.part.rels.values():
        if rel.reltype == comments_rel_type:
            comments_xml = rel.target_part
            break

    ts = datetime.utcnow().strftime("%Y-%m-%dT%H:%M:%SZ")

    if comments_xml is None:
        # Create new comments.xml
        from lxml import etree

        nsmap = {
            "w": W,
            "w14": W14,
            "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
            "mc": "http://schemas.openxmlformats.org/markup-compatibility/2006",
        }
        comments_root = etree.Element(f"{{{W}}}comments", nsmap=nsmap)

        comment_elem = etree.SubElement(comments_root, f"{{{W}}}comment")
        comment_elem.set(f"{{{W}}}id", str(comment_id))
        comment_elem.set(f"{{{W}}}author", author)
        comment_elem.set(f"{{{W}}}date", ts)
        comment_elem.set(f"{{{W}}}initials", initials)

        p = etree.SubElement(comment_elem, f"{{{W}}}p")
        p.set(f"{{{W14}}}paraId", f"{comment_id:08X}")
        p.set(f"{{{W14}}}textId", "77777777")

        # Annotation ref run
        ref_run = etree.SubElement(p, f"{{{W}}}r")
        ref_rpr = etree.SubElement(ref_run, f"{{{W}}}rPr")
        ref_style = etree.SubElement(ref_rpr, f"{{{W}}}rStyle")
        ref_style.set(f"{{{W}}}val", "CommentReference")
        etree.SubElement(ref_run, f"{{{W}}}annotationRef")

        # Text run
        text_run = etree.SubElement(p, f"{{{W}}}r")
        t = etree.SubElement(text_run, f"{{{W}}}t")
        t.text = text

        xml_bytes = etree.tostring(comments_root, xml_declaration=True, encoding="UTF-8")

        # Add as a new part
        from docx.opc.part import Part
        from docx.opc.packuri import PackURI

        part = Part(
            PackURI("/word/comments.xml"),
            "application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml",
            xml_bytes,
            doc.part.package,
        )
        doc.part.relate_to(part, comments_rel_type)
    else:
        # Append to existing comments.xml
        from lxml import etree

        root = etree.fromstring(comments_xml.blob)

        comment_elem = etree.SubElement(root, f"{{{W}}}comment")
        comment_elem.set(f"{{{W}}}id", str(comment_id))
        comment_elem.set(f"{{{W}}}author", author)
        comment_elem.set(f"{{{W}}}date", ts)
        comment_elem.set(f"{{{W}}}initials", initials)

        p = etree.SubElement(comment_elem, f"{{{W}}}p")
        p.set(f"{{{W14}}}paraId", f"{comment_id:08X}")
        p.set(f"{{{W14}}}textId", "77777777")

        ref_run = etree.SubElement(p, f"{{{W}}}r")
        ref_rpr = etree.SubElement(ref_run, f"{{{W}}}rPr")
        ref_style = etree.SubElement(ref_rpr, f"{{{W}}}rStyle")
        ref_style.set(f"{{{W}}}val", "CommentReference")
        etree.SubElement(ref_run, f"{{{W}}}annotationRef")

        text_run = etree.SubElement(p, f"{{{W}}}r")
        t = etree.SubElement(text_run, f"{{{W}}}t")
        t.text = text

        comments_xml._blob = etree.tostring(root, xml_declaration=True, encoding="UTF-8")


def create_early_revision(output_path: Path):
    """Create the early revision document with comments."""
    doc = Document()
    doc._comment_id = 0

    # Title
    title = doc.add_heading("Project Alpha - Technical Specification", level=0)

    # Section 1
    h1 = doc.add_heading("1. Introduction", level=1)
    p1 = doc.add_paragraph(
        "This document describes the technical specification for Project Alpha. "
        "The system will provide a cloud-based platform for data analytics "
        "with support for real-time processing and batch operations."
    )
    add_comment(doc, p1,
                "Should we specify which cloud providers are supported?",
                author="Alice Johnson", initials="AJ")

    p2 = doc.add_paragraph(
        "The target deployment date is Q3 2025, with beta testing beginning "
        "in Q2 2025."
    )
    add_comment(doc, p2,
                "This timeline seems aggressive. Can we get PM to confirm?",
                author="Bob Smith", initials="BS")

    # Section 2
    h2 = doc.add_heading("2. System Architecture", level=1)
    p3 = doc.add_paragraph(
        "The system uses a microservices architecture deployed on Kubernetes. "
        "Each service communicates via gRPC with a REST API gateway for "
        "external consumers."
    )
    add_comment(doc, p3,
                "We should add a diagram here showing the service topology.",
                author="Alice Johnson", initials="AJ")

    doc.add_heading("2.1 Data Layer", level=2)
    p4 = doc.add_paragraph(
        "The primary data store is PostgreSQL 14 with read replicas. "
        "Redis is used for caching and session management. "
        "Event streaming uses Apache Kafka."
    )
    add_comment(doc, p4,
                "Should we consider upgrading to PostgreSQL 16 for the new features?",
                author="Charlie Dev", initials="CD")

    doc.add_heading("2.2 Processing Engine", level=2)
    p5 = doc.add_paragraph(
        "Batch processing is handled by Apache Spark clusters. "
        "Real-time processing uses Apache Flink with exactly-once semantics."
    )

    # Section 3 - Table
    doc.add_heading("3. API Endpoints", level=1)
    p6 = doc.add_paragraph(
        "The following table lists the primary API endpoints for the system."
    )

    table = doc.add_table(rows=4, cols=4)
    table.style = "Table Grid"
    headers = ["Endpoint", "Method", "Description", "Auth Required"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h

    data = [
        ["/api/v1/data", "GET", "Retrieve data records", "Yes"],
        ["/api/v1/data", "POST", "Create new data record", "Yes"],
        ["/api/v1/health", "GET", "Health check endpoint", "No"],
    ]
    for row_idx, row_data in enumerate(data):
        for col_idx, val in enumerate(row_data):
            table.rows[row_idx + 1].cells[col_idx].text = val

    # Section 4
    doc.add_heading("4. Security Requirements", level=1)
    p7 = doc.add_paragraph(
        "All API endpoints must use TLS 1.2 or higher. "
        "Authentication is handled via OAuth 2.0 with JWT tokens. "
        "Tokens expire after 30 minutes."
    )
    add_comment(doc, p7,
                "Legal team requires TLS 1.3 minimum. Please update.",
                author="Diana Security", initials="DS")

    p8 = doc.add_paragraph(
        "Data at rest must be encrypted using AES-256. "
        "Key management uses AWS KMS."
    )

    # Section 5
    doc.add_heading("5. Performance Targets", level=1)
    bullets = [
        "API response time: < 200ms (p95)",
        "Data ingestion throughput: 10,000 records/second",
        "System uptime: 99.9%",
        "Recovery time objective (RTO): 4 hours",
        "Recovery point objective (RPO): 1 hour",
    ]
    for b in bullets:
        doc.add_paragraph(b, style="List Bullet")

    p9 = doc.add_paragraph(
        "These targets apply to the production environment under normal load conditions."
    )
    add_comment(doc, p9,
                "We need to define what 'normal load' means quantitatively.",
                author="Bob Smith", initials="BS")

    doc.save(str(output_path))
    print(f"Created early revision: {output_path}")


def create_latest_revision(output_path: Path):
    """Create the latest revision with edits (no comments)."""
    doc = Document()

    # Title (unchanged)
    doc.add_heading("Project Alpha - Technical Specification", level=0)

    # Section 1 (modified)
    doc.add_heading("1. Introduction", level=1)
    doc.add_paragraph(
        "This document describes the technical specification for Project Alpha. "
        "The system will provide a multi-cloud platform for data analytics "
        "with support for real-time processing, batch operations, and ML inference."
    )
    doc.add_paragraph(
        "The target deployment date is Q4 2025, with beta testing beginning "
        "in Q3 2025. The timeline has been adjusted per stakeholder feedback."
    )

    # Section 2 (modified)
    doc.add_heading("2. System Architecture", level=1)
    doc.add_paragraph(
        "The system uses a microservices architecture deployed on Kubernetes (EKS/GKE). "
        "Each service communicates via gRPC with a REST API gateway for "
        "external consumers. Service mesh is provided by Istio."
    )

    doc.add_heading("2.1 Data Layer", level=2)
    doc.add_paragraph(
        "The primary data store is PostgreSQL 16 with read replicas and "
        "connection pooling via PgBouncer. "
        "Redis Cluster is used for caching and session management. "
        "Event streaming uses Apache Kafka with Schema Registry."
    )

    doc.add_heading("2.2 Processing Engine", level=2)
    doc.add_paragraph(
        "Batch processing is handled by Apache Spark 3.5 clusters with "
        "dynamic resource allocation. "
        "Real-time processing uses Apache Flink with exactly-once semantics "
        "and checkpointing every 30 seconds."
    )

    # NEW section
    doc.add_heading("2.3 ML Pipeline", level=2)
    doc.add_paragraph(
        "Machine learning workloads run on dedicated GPU nodes using "
        "Kubeflow for orchestration. Model serving uses TensorFlow Serving "
        "and Triton Inference Server."
    )

    # Section 3 - Table (modified)
    doc.add_heading("3. API Endpoints", level=1)
    doc.add_paragraph(
        "The following table lists the primary API endpoints for the system. "
        "All endpoints follow the OpenAPI 3.0 specification."
    )

    table = doc.add_table(rows=6, cols=4)
    table.style = "Table Grid"
    headers = ["Endpoint", "Method", "Description", "Auth Required"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h

    data = [
        ["/api/v2/data", "GET", "Retrieve data records with pagination", "Yes"],
        ["/api/v2/data", "POST", "Create new data record", "Yes"],
        ["/api/v2/data/{id}", "PUT", "Update existing record", "Yes"],
        ["/api/v2/models", "POST", "Submit ML inference request", "Yes"],
        ["/api/v2/health", "GET", "Health check endpoint", "No"],
    ]
    for row_idx, row_data in enumerate(data):
        for col_idx, val in enumerate(row_data):
            table.rows[row_idx + 1].cells[col_idx].text = val

    # Section 4 (modified)
    doc.add_heading("4. Security Requirements", level=1)
    doc.add_paragraph(
        "All API endpoints must use TLS 1.3 or higher. "
        "Authentication is handled via OAuth 2.0 with JWT tokens. "
        "Tokens expire after 15 minutes with refresh token support."
    )
    doc.add_paragraph(
        "Data at rest must be encrypted using AES-256-GCM. "
        "Key management uses cloud-native KMS (AWS KMS or GCP Cloud KMS). "
        "Key rotation occurs every 90 days."
    )

    # Section 5 (modified)
    doc.add_heading("5. Performance Targets", level=1)
    bullets = [
        "API response time: < 150ms (p95), < 500ms (p99)",
        "Data ingestion throughput: 50,000 records/second",
        "ML inference latency: < 100ms (p95)",
        "System uptime: 99.95%",
        "Recovery time objective (RTO): 1 hour",
        "Recovery point objective (RPO): 15 minutes",
    ]
    for b in bullets:
        doc.add_paragraph(b, style="List Bullet")

    doc.add_paragraph(
        "These targets apply to the production environment under normal load "
        "conditions, defined as up to 1,000 concurrent users and 100 requests/second."
    )

    # NEW Section 6
    doc.add_heading("6. Monitoring and Observability", level=1)
    doc.add_paragraph(
        "The system uses a comprehensive observability stack including "
        "Prometheus for metrics, Grafana for dashboards, and Jaeger for "
        "distributed tracing. All services emit structured logs in JSON format."
    )

    doc.save(str(output_path))
    print(f"Created latest revision: {output_path}")


def main():
    test_dir = Path("test_documents")
    test_dir.mkdir(exist_ok=True)

    create_early_revision(test_dir / "test_early_rev.docx")
    create_latest_revision(test_dir / "test_latest_rev.docx")

    print(f"\nTest documents created in {test_dir}/")
    print("\nTo run the comparison:")
    print(f"  python compare_revisions.py "
          f"{test_dir}/test_early_rev.docx "
          f"{test_dir}/test_latest_rev.docx "
          f"output_comparison.docx --verbose")


if __name__ == "__main__":
    main()
